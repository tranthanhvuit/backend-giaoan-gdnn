from fastapi import FastAPI, UploadFile, File, Form, HTTPException
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import StreamingResponse
from dotenv import load_dotenv
import os
from io import BytesIO
from docx import Document
from openai import OpenAI
from pathlib import Path

# Load biến môi trường
load_dotenv()
client = OpenAI(api_key=os.getenv("OPENAI_API_KEY"))

app = FastAPI()

@app.get("/")
def root():
    return {"message": "Backend is running"}

# CORS
app.add_middleware(
    CORSMiddleware,
    allow_origins=[
        "https://chuyen-decuong-giaoan.onrender.com",
        "http://localhost:3000"
    ],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

# Load prompt mẫu từ thư mục
PROMPT_TEMPLATES = {
    "lythuyet": Path("prompts/prompt_giaoan_lythuyet.txt").read_text(encoding="utf-8"),
    "thuchanh": Path("prompts/prompt_giaoan_thuchanh.txt").read_text(encoding="utf-8"),
    "tichhop": Path("prompts/prompt_giaoan_tichhop.txt").read_text(encoding="utf-8"),
}

def read_docx(file_bytes):
    doc = Document(BytesIO(file_bytes))
    full_text = [para.text for para in doc.paragraphs]
    return "\n".join(full_text)

def generate_giao_an(decuong_text, loai):
    prompt_template = PROMPT_TEMPLATES.get(loai, "")
    full_prompt = f"{prompt_template}\n\n{decuong_text}"

    response = client.chat.completions.create(
        model="gpt-4",
        messages=[{"role": "user", "content": full_prompt}],
        temperature=0.3,
        max_tokens=2000,
    )
    return response.choices[0].message.content

@app.post("/generate")
async def generate(file: UploadFile = File(...), loai: str = Form(...)):
    content = await file.read()
    decuong_text = read_docx(content)
    result = generate_giao_an(decuong_text, loai)
    return {"result": result}

@app.post("/generate-docx")
async def generate_docx(file: UploadFile = File(...), loai: str = Form(...)):
    try:
        contents = await file.read()
        if not contents:
            raise HTTPException(status_code=400, detail="Tệp trống")

        decuong_text = read_docx(contents)
        giaoan_text = generate_giao_an(decuong_text, loai)

        # Tạo file Word từ nội dung GPT
        doc = Document()
        doc.add_paragraph(giaoan_text)

        doc_stream = BytesIO()
        doc.save(doc_stream)
        doc_stream.seek(0)

        return StreamingResponse(
            doc_stream,
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            headers={"Content-Disposition": "attachment; filename=giao_an.docx"}
        )

    except Exception as e:
        print("Lỗi khi xử lý file:", e)
        raise HTTPException(status_code=500, detail=f"Lỗi xử lý: {str(e)}")